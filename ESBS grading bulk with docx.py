
import os
import docx
from openai import OpenAI
from typing import Dict, List
import json
from PyQt6.QtWidgets import QApplication, QFileDialog
import sys

class DocumentGrader:
    def __init__(self, api_key: str):
        """Initialize the grader with OpenAI API key"""
        self.client = OpenAI(api_key=api_key)
        
    def read_docx(self, file_path: str) -> str:
        """Read content from a Word document"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            return '\n'.join(full_text)
        except Exception as e:
            raise Exception(f"Error reading document {file_path}: {str(e)}")
    
    def prepare_grading_prompt(self, exercise: str, rubric: str, solution: str) -> str:
        """Prepare the prompt for OpenAI"""
        prompt = f'''Please grade the following student solution based on the exercise description and grading rubric provided.

    Exercise Description:
    {exercise}

    Grading Rubric:
    {rubric}

    Students Solution:
    {solution}

    Please provide:
    1. A detailed analysis of the solution
    2. Points awarded for each rubric criterion
    3. Total score
    4. Specific feedback and suggestions for improvement

    Format your response ONLY as JSON with the following structure:
    {{
        "analysis": "detailed analysis here",
        "criterion_scores": {{
            "criterion1": "score",
            "criterion2": "score"
    }},
    "total_score": "numerical_score",
    "feedback": "detailed feedback here"
    }}
''' 
        return prompt
    
    def grade_solution(self, exercise_path: str, rubric_path: str, solution_path: str) -> Dict:
        """Grade the solution using OpenAI API"""
        try:
            # Read all documents
            exercise_text = self.read_docx(exercise_path)
            rubric_text = self.read_docx(rubric_path)
            solution_text = self.read_docx(solution_path)

            # Prepare the prompt
            prompt = self.prepare_grading_prompt(exercise_text, rubric_text, solution_text)

            # Call OpenAI API
            response = self.client.chat.completions.create(
                model="gpt-4-turbo-preview",  # or "gpt-4" or "gpt-3.5-turbo" depending on your needs
                messages=[
                    {"role": "system", "content": "You are an expert grader who evaluates student solutions accurately and fairly."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.0,
                max_tokens=4000,
                response_format={"type": "json_object"}  # Add this line to force JSON response
            )

            response_content= response.choices[0].message.content.strip()
            print("Raw API Rsponse:", response_content)

            # Parse the response
            try:
                grading_result = json.loads(response_content)
                return grading_result
            except json.JSONDecodeError:
                raise Exception("Failed to parse API response as JSON")

        except Exception as e:
            raise Exception(f"Grading failed: {str(e)}")
        
    def save_results(self, results: Dict, output_path: str):
        """Save grading results to a Word document"""
        try:
            doc = docx.Document()
            
            # Add results to document
            doc.add_heading('Grading Results', 0)
        
            # Add Analysis section
            doc.add_heading('Analysis', level=1)
            doc.add_paragraph(str(results.get('analysis', 'No analysis provided')))

            # Add criterio scores section    
            doc.add_heading('Criterion Scores', level=1)
            criterion_scores=results.get('criterion_scores', {})
            if isinstance(criterion_scores, dict):
                for criterion, score in criterion_scores.items():
                    doc.add_paragraph(f"{criterion}: {score}")
            else:
                doc.add_paragraph("No criterion scores available")
            
            # Add Total score section
            doc.add_heading('Total Score', level=1)
            doc.add_paragraph(str(results.get('total_score', 'No score available')))
            
            # Add Feedback section
            doc.add_heading('Feedback', level=1)
            doc.add_paragraph(str(results.get('feedback','No feeback provided')))
            
            # Save the document
            doc.save(output_path)
            print(f"Results succesfully saver to {output_path}")

        except Exception as e:
            raise Exception(f"Error saving results: {str(e)}")
        
    def select_filename(title="Select file"):
        """
        Opens a file dialog to select a document
        Returns the selected file path
        """
        app = QApplication(sys.argv)
        # Create native macOS file dialog
        file_dialog = QFileDialog()
        file_dialog.setFileMode(QFileDialog.FileMode.ExistingFile)
        file_dialog.setNameFilter("Word documents (*.docx);;All Files (*)")
        file_dialog.setWindowTitle(title)

        # Set dialog to use narive MacOS style
        file_dialog.setOption(QFileDialog.DialogCode.Option.DontUseNativeDialog, False)

        if file_dialog.exec() == QFileDialog.DialogCode.Accepted:
            selected_files = file_dialog.selectedFiles()
            if selected_files:
                return selected_files[0]
            return None

    def create_output_filename(input_path: str) -> str:
        """
        Creates an output forlename based on the input filename by adding '_graded' before the extension.
        """
        directory = os.path.dirname(input_path)
        filename = os.path.basename(input_path)
        name, ext = os.path.splitext(filename)
        return os.path.joindirectory(directory, f"{name}_graded{ext}")




def main():
    try:
        # Get API key from environment variable
        MI_CLAVE = os.getenv('MI_CLAVE_API_OPENAI')
        if not MI_CLAVE:
            raise ValueError("La clave de la API de OpenAI no se encuentra en las variables de entorno")

        # Initialize grader
        grader = DocumentGrader(MI_CLAVE)

        # Define file paths
        exercise_path = "/Users/JJVR/Library/CloudStorage/OneDrive-KALEIDAGEOGRAFIAS&MERCADOSSL/ESBS/Asignaturas/Plagiarism & AI/Exercise/exercise instructions.docx"  # Path to exercise description
        rubric_path = "/Users/JJVR/Library/CloudStorage/OneDrive-KALEIDAGEOGRAFIAS&MERCADOSSL/ESBS/Asignaturas/Plagiarism & AI/Exercise/grading rubric.docx"     # Path to grading rubric
        solution_path = "/Users/JJVR/Library/CloudStorage/OneDrive-KALEIDAGEOGRAFIAS&MERCADOSSL/ESBS/Asignaturas/Plagiarism & AI/Class exercise/Plagiarism mma Cy Carre.docx"  # Path to student's solution
        output_path = "/Users/JJVR/Library/CloudStorage/OneDrive-KALEIDAGEOGRAFIAS&MERCADOSSL/ESBS/Asignaturas/Plagiarism & AI/Class exercise/grading_results.docx"  # Path for output file

        # Check if input files exist
        for path in [exercise_path, rubric_path, solution_path]:
            if not os.path.exists(path):
                raise FileNotFoundError(f"Input file not found: {path}")

        # Grade the solution
        results = grader.grade_solution(exercise_path, rubric_path, solution_path)

        # Save the results
        grader.save_results(results, output_path)

        print(f"Grading completed successfully. Results saved to {output_path}")

    except Exception as e:
        print(f"Error: {str(e)}")
        exit(1)

if __name__ == "__main__":
    main()