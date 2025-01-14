import os
import docx
import anthropic
from typing import Dict, List
import json

class DocumentGrader:
    def __init__(self, api_key: str):
        """Initialize the grader with Anthropic API key"""
        self.client = anthropic.Client(api_key)
        
    def read_docx(self, file_path: str) -> str:
        """Read content from a Word document"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            return '\n'.join(full_text)
        except Exception as e:
            raise Exception(f"Error reading document {file_path}: {str(e)}")

    def prepare_grading_prompt(self, exercise: str, rubric: str, solution: str) -> str:
        """Prepare the prompt for Anthropic"""
        prompt = f"""Please grade the following student solution based on the exercise description and grading rubric provided.

Exercise Description:
{exercise}

Grading Rubric:
{rubric}

Student's Solution:
{solution}

Please provide:
1. A detailed analysis of the solution
2. Points awarded for each rubric criterion
3. Total score
4. Specific feedback and suggestions for improvement

Format your response as JSON with the following structure:
{{
    "analysis": "detailed analysis here",
    "criterion_scores": {{
        "criterion1": "score",
        "criterion2": "score"
    }},
    "total_score": "final_score",
    "feedback": "detailed feedback here"
}}
"""
        return prompt

    def grade_solution(self, exercise_path: str, rubric_path: str, solution_path: str) -> Dict:
        """Grade the solution using Anthropic API"""
        try:
            # Read all documents
            exercise_text = self.read_docx(exercise_path)
            rubric_text = self.read_docx(rubric_path)
            solution_text = self.read_docx(solution_path)

            # Prepare the prompt
            prompt = self.prepare_grading_prompt(exercise_text, rubric_text, solution_text)

            # Call Anthropic API
            response = self.client.messages.create(
                model="claude-3-opus-20240229",
                max_tokens=4000,
                temperature=0.0,
                messages=[{
                    "role": "user",
                    "content": prompt
                }]
            )

            # Parse the response
            try:
                grading_result = json.loads(response.content)
                return grading_result
            except json.JSONDecodeError:
                raise Exception("Failed to parse API response as JSON")

        except Exception as e:
            raise Exception(f"Grading failed: {str(e)}")

    def save_results(self, results: Dict, output_path: str):
        """Save grading results to a Word document"""
        try:
            doc = docx.Document()
            
            # Add results to document
            doc.add_heading('Grading Results', 0)
            
            doc.add_heading('Analysis', level=1)
            doc.add_paragraph(results['analysis'])
            
            doc.add_heading('Criterion Scores', level=1)
            for criterion, score in results['criterion_scores'].items():
                doc.add_paragraph(f"{criterion}: {score}")
            
            doc.add_heading('Total Score', level=1)
            doc.add_paragraph(str(results['total_score']))
            
            doc.add_heading('Feedback', level=1)
            doc.add_paragraph(results['feedback'])
            
            doc.save(output_path)
        except Exception as e:
            raise Exception(f"Error saving results: {str(e)}")

def main():
    # Get API key from environment variable
   api_key = os.getenv("MI_CLAVE_API_ANTROPIC")