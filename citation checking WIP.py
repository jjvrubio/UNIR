
import re
import json
import Cocoa
import datetime
from Quartz import PDFDocument
from Foundation import NSURL
from docx import Documentimport os
import re
import json
import Cocoa
import datetime
from Quartz import PDFDocument
from Foundation import NSURL
from docx import Document


def debug_log(message, output_dir="", title="debug_log"):
    """Appends debug messages to a Markdown log file in the specified directory."""
    if not output_dir:
        output_dir = os.getcwd()  # Default to the current working directory

    debug_file = os.path.join(output_dir, f"{title}.md")
    try:
        with open(debug_file, "a") as log_file:
            timestamp = datetime.datetime.now().isoformat()
            log_file.write(f"- **[{timestamp}]** {message}\n")
    except Exception as e:
        print(f"Failed to write debug log: {e}")


def choose_file(file_type):
    """Use a native macOS file dialog to select a file and return its path and directory."""
    panel = Cocoa.NSOpenPanel.openPanel()
    panel.setAllowsMultipleSelection_(False)
    panel.setCanChooseDirectories_(False)

    if file_type == "thesis":
        panel.setAllowedFileTypes_(["pdf", "docx"])
    elif file_type == "references":
        panel.setAllowedFileTypes_(["json"])

    if panel.runModal() == Cocoa.NSModalResponseOK:
        file_path = panel.URLs()[0].path()
        file_dir = os.path.dirname(file_path)
        debug_log(f"Selected {file_type} file: {file_path}", file_dir, "debug_log")
        return file_path, file_dir
    else:
        debug_log(f"No {file_type} file selected.", "", "debug_log")
        return None, None


def get_file_title(file_path):
    """Extract the file title (name without extension) from the file path."""
    return os.path.splitext(os.path.basename(file_path))[0]


def extract_text(file_path, source_type):
    """Generalized text extraction for PDF and DOCX formats."""
    if source_type == "pdf":
        return extract_pdf(file_path)
    elif source_type == "docx":
        return extract_docx(file_path)
    else:
        raise ValueError("Unsupported file format")


def extract_pdf(file_path):
    """Extract raw text from a PDF."""
    url = NSURL.fileURLWithPath_(file_path)
    pdf_doc = PDFDocument.alloc().initWithURL_(url)
    text = ""

    debug_log(f"Starting PDF extraction for file: {file_path}")

    for page_num in range(pdf_doc.pageCount()):
        page = pdf_doc.pageAtIndex_(page_num)
        page_text = page.string()
        if page_text:
            debug_log(f"Extracted text from page {page_num + 1}.")
        else:
            debug_log(f"No text found on page {page_num + 1}.")
        text += page_text

    debug_log(f"PDF extraction completed. Total length: {len(text)} characters.")
    return text


def extract_docx(file_path):
    """Extract raw text from a DOCX file."""
    doc = Document(file_path)
    text = ""

    debug_log(f"Starting DOCX extraction for file: {file_path}")

    for i, paragraph in enumerate(doc.paragraphs, start=1):
        if paragraph.text.strip():
            debug_log(f"Extracted text from paragraph {i}.")
        text += paragraph.text + "\n"

    debug_log(f"DOCX extraction completed. Total length: {len(text)} characters.")
    return text


def process_extracted_text(raw_text):
    """Process raw extracted text, stopping at 'Anexo' and skipping 'Índice'."""
    debug_log("Starting text processing.")

    if not raw_text.strip():
        debug_log("Raw text is empty. Processing aborted.")
        return ""

    lines = raw_text.split("\n")
    processed_text = []
    found_toc = False

    for line in lines:
        if not found_toc and ("Índice" in line or "CONTENIDO" in line.upper()):
            debug_log("Detected Table of Contents.")
            found_toc = True
            continue

        if found_toc and ("Anexo" in line or "ANEXO" in line):
            debug_log("Detected 'Anexo' section. Stopping processing.")
            break

        processed_text.append(line)

    debug_log(f"Processed text length: {len(processed_text)} lines.")
    return "\n".join(processed_text)


def load_references(json_path):
    """Load references from JSON and parse relevant details."""
    try:
        with open(json_path, "r") as file:
            references = json.load(file)
    except json.JSONDecodeError as e:
        debug_log(f"Error loading JSON file: {e}")
        return []

    debug_log(f"Loaded {len(references)} references.")

    parsed_references = []
    for ref in references:
        match = re.search(r"(.*?)\.\s\((\d{4})\)\.\s(.*?)\.", ref["reference"])
        if match:
            parsed_references.append({
                "author": match.group(1),
                "year": match.group(2),
                "title": match.group(3),
                "reference": ref["reference"],
                "is_compliant": ref["is_compliant"],
                "url": ref.get("url", False)
            })
        else:
            debug_log(f"Failed to parse reference: {ref['reference']}")

    debug_log(f"Parsed {len(parsed_references)} valid references.")
    return parsed_references


def search_citations_in_text(full_text):
    """Locate and log citations in the text, capturing their locations."""
    debug_log("Starting citation search in extracted text.")

    citations_with_locations = []
    lines = full_text.split("\n")  # Split into lines for paragraph-like processing

    for index, line in enumerate(lines, start=1):
        matches = re.findall(r"\(([^,]+), (\d{4})(?:, p. \d+)?\)", line)
        for match in matches:
            citations_with_locations.append({"citation": match, "location": index})
            debug_log(f"Found citation: {match} at line {index}")

    debug_log(f"Total citations found: {len(citations_with_locations)}")
    return citations_with_locations


def generate_report_with_locations(matched, unmatched, references, output_dir, title, report_name="citation_report.md"):
    """Generate a markdown report with citation locations and save it in the specified directory."""
    report_name = f"{title}_{report_name}"
    output_path = os.path.join(output_dir, report_name)

    cited_references = {f"{c['citation'][0]}, {c['citation'][1]}" for c in matched}
    uncited_references = [
        ref for ref in references
        if f"{ref['author']}, {ref['year']}" not in cited_references
    ]

    debug_log(f"Generating Markdown report at: {output_path}", output_dir, title)

    with open(output_path, "w") as report:
        report.write("# Citation Analysis Report\n\n")
        report.write("## Summary\n")
        report.write(f"- Total Citations: {len(matched) + len(unmatched)}\n")
        report.write(f"- Properly Cited: {len(matched)}\n")
        report.write(f"- Improperly Cited: {len(unmatched)}\n")
        report.write(f"- References Not Cited: {len(uncited_references)}\n\n")

        report.write("## Detailed Citations:\n")
        if matched:
            report.write("\n### Properly Cited:\n")
            for item in matched:
                location = item.get("location", "Unknown")
                report.write(f"- {item['citation'][0]}, {item['citation'][1]} (Location: {location})\n")
        
        if unmatched:
            report.write("\n### Improperly Cited:\n")
            for item in unmatched:
                location = item.get("location", "Unknown")
                report.write(f"- {item['citation'][0]}, {item['citation'][1]} (Location: {location})\n")

        report.write("\n## References Not Cited:\n")
        for ref in uncited_references:
            report.write(f"- {ref['author']} ({ref['year']}). {ref['title']}\n")

    debug_log(f"Markdown report successfully saved at: {output_path}", output_dir, title)



def main():
    thesis_file, thesis_dir = choose_file("thesis")
    references_file, _ = choose_file("references")

    if not thesis_file or not references_file:
        debug_log("File selection aborted.", thesis_dir or os.getcwd(), "debug_log")
        return

    # Extract the title from the thesis filename
    thesis_title = get_file_title(thesis_file)

    debug_log("Program started.", thesis_dir, thesis_title)

    source_type = "pdf" if thesis_file.endswith(".pdf") else "docx"
    raw_text = extract_text(thesis_file, source_type)
    debug_log(f"Extracted raw text length: {len(raw_text)} characters.", thesis_dir, thesis_title)

    filtered_text = process_extracted_text(raw_text)
    debug_log(f"Filtered text length: {len(filtered_text)} characters.", thesis_dir, thesis_title)

    references = load_references(references_file)
    citations_with_locations = search_citations_in_text(filtered_text)

    # Match citations with references
    matched = []
    unmatched = []
    for citation in citations_with_locations:
        author, year = citation["citation"]
        if any(ref for ref in references if ref["author"] == author and ref["year"] == year):
            matched.append(citation)
        else:
            unmatched.append(citation)

    generate_report_with_locations(matched, unmatched, references, thesis_dir, thesis_title)
    debug_log("Program completed successfully.", thesis_dir, thesis_title)



if __name__ == "__main__":
    main()




def debug_log(message):
    """Appends debug messages to a Markdown log file."""
    with open("debug_log.md", "a") as log_file:
        timestamp = datetime.datetime.now().isoformat()
        log_file.write(f"- **[{timestamp}]** {message}\n")



def choose_file(file_type):
    """Use a native macOS file dialog to select a file."""
    panel = Cocoa.NSOpenPanel.openPanel()
    panel.setAllowsMultipleSelection_(False)
    panel.setCanChooseDirectories_(False)

    if file_type == "thesis":
        panel.setAllowedFileTypes_(["pdf", "docx"])
    elif file_type == "references":
        panel.setAllowedFileTypes_(["json"])

    if panel.runModal() == Cocoa.NSModalResponseOK:
        file_path = panel.URLs()[0].path()
        debug_log(f"Selected {file_type} file: {file_path}")
        return file_path
    else:
        debug_log(f"No {file_type} file selected.")
        return None



def extract_pdf_text(file_path):
    """Extract text from a PDF while excluding annexes and considering the table of contents."""
    url = NSURL.fileURLWithPath_(file_path)
    pdf_doc = PDFDocument.alloc().initWithURL_(url)
    text = ""
    is_after_toc = False

    debug_log("Starting PDF text extraction.")

    for page_num in range(pdf_doc.pageCount()):
        page = pdf_doc.pageAtIndex_(page_num)
        page_text = page.string()

        if not is_after_toc and "Índice" in page_text:
            debug_log(f"Table of Contents detected on page {page_num + 1}.")
            is_after_toc = True
            continue

        if "Anexo" in page_text:
            debug_log(f"Anexo section detected on page {page_num + 1}. Stopping extraction.")
            break

        if is_after_toc:
            text += page_text

    debug_log("PDF text extraction completed.")
    return text


def extract_docx_text(file_path):
    """Extract text from a DOCX file while excluding annexes and considering the table of contents."""
    doc = Document(file_path)
    text = ""
    is_after_toc = False

    debug_log("Starting DOCX text extraction.")

    for paragraph in doc.paragraphs:
        if not is_after_toc and "Índice" in paragraph.text:
            debug_log("Table of Contents detected.")
            is_after_toc = True
            continue

        if "Anexo" in paragraph.text:
            debug_log("Anexo section detected. Stopping extraction.")
            break

        if is_after_toc:
            text += paragraph.text + "\n"

    debug_log("DOCX text extraction completed.")
    return text

def extract_text_until_anexo(full_text):
    """
    Process the extracted text and stop at the first actual occurrence of 'Anexo' 
    while ignoring mentions in the Table of Contents.
    """
    debug_log("Processing text to stop at the first actual 'Anexo' section.")

    # Split the text into lines for processing
    lines = full_text.split("\n")
    processed_text = []
    found_toc = False

    for line in lines:
        # Detect potential Table of Contents
        if not found_toc and ("Índice" in line or "CONTENIDO" in line.upper()):
            debug_log("Detected Table of Contents.")
            found_toc = True
            continue

        # Stop processing at the first real "Anexo" section
        if found_toc and ("Anexo" in line or "ANEXO" in line):
            debug_log("Detected the start of the 'Anexo' section. Stopping.")
            break

        processed_text.append(line)

    debug_log("Finished processing text until the 'Anexo' section.")
    return "\n".join(processed_text)


def search_citations_in_text(full_text):
    """
    Locate and log citations in the text, capturing their locations.
    For PDF: Use page numbers.
    For DOCX: Use paragraph numbers.
    """
    debug_log("Starting citation search in extracted text.")

    citations_with_locations = []
    lines = full_text.split("\n")  # Split into lines for paragraph-like processing

    for index, line in enumerate(lines, start=1):
        matches = re.findall(r"\(([^,]+), (\d{4})(?:, p. \d+)?\)", line)
        for match in matches:
            citations_with_locations.append({"citation": match, "location": index})
            debug_log(f"Found citation: {match} at line {index}")

    debug_log(f"Total citations found: {len(citations_with_locations)}")
    return citations_with_locations


def load_references(json_path):
    """Load references from JSON and parse relevant details."""
    with open(json_path, "r") as file:
        references = json.load(file)

    debug_log(f"Loaded {len(references)} references.")

    parsed_references = []
    for ref in references:
        match = re.search(r"(.*?)\.\s\((\d{4})\)\.\s(.*?)\.", ref["reference"])
        if match:
            parsed_references.append({
                "author": match.group(1),
                "year": match.group(2),
                "title": match.group(3),
                "reference": ref["reference"],
                "is_compliant": ref["is_compliant"],
                "url": ref.get("url", False)
            })
        else:
            debug_log(f"Failed to parse reference: {ref['reference']}")

    debug_log(f"Parsed {len(parsed_references)} valid references.")
    return parsed_references

def validate_citations_with_locations(thesis_text, references, source_type="pdf"):
    """Validate citations and include their locations in the text."""
    citations_with_locations = []
    unmatched_citations = []

    # For PDF: track page numbers; for DOCX: approximate paragraph or block location
    if source_type == "pdf":
        pages = thesis_text.split("\f")  # Split text into pages by form-feed character
        for page_num, page_text in enumerate(pages, start=1):
            matches = re.findall(r"\(([^,]+), (\d{4})(?:, p. \d+)?\)", page_text)
            for match in matches:
                citations_with_locations.append({"citation": match, "page": page_num})
    elif source_type == "docx":
        paragraphs = thesis_text.split("\n")  # Split text into paragraphs
        for paragraph_num, paragraph_text in enumerate(paragraphs, start=1):
            matches = re.findall(r"\(([^,]+), (\d{4})(?:, p. \d+)?\)", paragraph_text)
            for match in matches:
                citations_with_locations.append({"citation": match, "paragraph": paragraph_num})

    # Match citations with references
    matched_citations = []
    for item in citations_with_locations:
        citation = item["citation"]
        author, year = citation
        if any(ref for ref in references if ref["author"] == author and ref["year"] == year):
            matched_citations.append(item)
        else:
            unmatched_citations.append(item)

    debug_log(f"Total citations found: {len(citations_with_locations)}")
    debug_log(f"Matched citations: {len(matched_citations)}")
    debug_log(f"Unmatched citations: {len(unmatched_citations)}")

    return matched_citations, unmatched_citations


def generate_report_with_locations(matched, unmatched, references, file_name="report.md"):
    """Generate a markdown report with citation locations and detailed compliance information."""
    cited_references = {f"{c['citation'][0]}, {c['citation'][1]}" for c in matched}
    uncited_references = [
        ref for ref in references
        if f"{ref['author']}, {ref['year']}" not in cited_references
    ]

    debug_log("Generating Markdown report with citation locations.")

    with open(file_name, "w") as report:
        report.write("# Citation Analysis Report\n\n")
        report.write("## Summary\n")
        report.write(f"- Total Citations: {len(matched) + len(unmatched)}\n")
        report.write(f"- Properly Cited: {len(matched)}\n")
        report.write(f"- Improperly Cited: {len(unmatched)}\n")
        report.write(f"- References Not Cited: {len(uncited_references)}\n\n")

        report.write("## Detailed Citations:\n")
        if matched:
            report.write("\n### Properly Cited:\n")
            for item in matched:
                location = item.get("page", item.get("paragraph", "Unknown"))
                report.write(f"- {item['citation'][0]}, {item['citation'][1]} (Location: {location})\n")
        
        if unmatched:
            report.write("\n### Improperly Cited:\n")
            for item in unmatched:
                location = item.get("page", item.get("paragraph", "Unknown"))
                report.write(f"- {item['citation'][0]}, {item['citation'][1]} (Location: {location})\n")

        report.write("\n## References Not Cited:\n")
        for ref in uncited_references:
            report.write(f"- {ref['author']} ({ref['year']}). {ref['title']}.\n")

    debug_log("Markdown report with locations generated successfully.")



def main():
    debug_log("Program started.")
    thesis_file = choose_file("thesis")
    references_file = choose_file("references")

    if not thesis_file or not references_file:
        debug_log("File selection aborted.")
        return

    # Extract the full text from the thesis
    if thesis_file.endswith(".pdf"):
        full_text = extract_pdf_text(thesis_file)
        source_type = "pdf"
    elif thesis_file.endswith(".docx"):
        full_text = extract_docx_text(thesis_file)
        source_type = "docx"
    else:
        debug_log("Unsupported thesis file format.")
        return

    # Process the text to stop at the first "Anexo"
    filtered_text = extract_text_until_anexo(full_text)

    # Load references
    references = load_references(references_file)

    # Locate citations in the filtered text
    citations_with_locations = search_citations_in_text(filtered_text)

    # Match citations with references
    matched = []
    unmatched = []
    for citation in citations_with_locations:
        author, year = citation["citation"]
        if any(ref for ref in references if ref["author"] == author and ref["year"] == year):
            matched.append(citation)
        else:
            unmatched.append(citation)

    # Generate the report
    generate_report_with_locations(matched, unmatched, references)
    debug_log("Program completed successfully.")
