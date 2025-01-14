#!/Users/JJVR/Documents/Personal/JJVR/automatizaciones/.venv_automatizar/bin/python

# This program has the right invocation ///from openai import OpenAI/// of Open AI. The calls are //// self.client.chat.completions. /////


import os
import PyPDF2
import docx
import markdown
from openai import OpenAI
from typing import Dict, List
import json
from Cocoa import NSOpenPanel, NSApplication, NSSavePanel
from Foundation import NSArray

def select_folder() -> str:
    """Select a folder containing student solutions"""
    NSApplication.sharedApplication()
    panel = NSOpenPanel.alloc().init()
    panel.setTitle_("Select Folder with Student Solutions")
    panel.setCanChooseFiles_(False)
    panel.setCanChooseDirectories_(True)
    panel.setAllowsMultipleSelection_(False)
    
    if panel.runModal() == 1:
        selected_folder = panel.URLs()[0].path()
        return selected_folder
    return None

def select_output_folder() -> str:
    """Select folder to save grading results"""
    NSApplication.sharedApplication()
    panel = NSOpenPanel.alloc().init()
    panel.setTitle_("Select Folder to Save Grading Results")
    panel.setCanChooseFiles_(False)
    panel.setCanChooseDirectories_(True)
    panel.setAllowsMultipleSelection_(False)
    
    if panel.runModal() == 1:
        output_folder = panel.URLs()[0].path()
        return output_folder
    return None

class DocumentGrader:
    def __init__(self, api_key: str):
        self.client = OpenAI(api_key=api_key)

    def read_pdf(self, file_path: str) -> str:
        try:
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                full_text = []
                for page in pdf_reader.pages:
                    full_text.append(page.extract_text())
                return '\n'.join(full_text)
        except Exception as e:
            raise Exception(f"Error reading PDF {file_path}: {str(e)}")

    def read_docx(self, file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            raise Exception(f"Error reading DOCX {file_path}: {str(e)}")

    def read_markdown(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as md_file:
                md_text = md_file.read()
                html = markdown.markdown(md_text)
                text = html.replace('<p>', '').replace('</p>', '\n')\
                          .replace('<h1>', '').replace('</h1>', '\n')\
                          .replace('<h2>', '').replace('</h2>', '\n')\
                          .replace('<h3>', '').replace('</h3>', '\n')
                return text
        except Exception as e:
            raise Exception(f"Error reading Markdown {file_path}: {str(e)}")

    def read_document(self, file_path: str) -> str:
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension == '.pdf':
            return self.read_pdf(file_path)
        elif file_extension == '.docx':
            return self.read_docx(file_path)
        elif file_extension in ['.md','.markdown']:
            return self.read_markdown(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def prepare_grading_prompt(self, exercise: str, rubric: str, solution: str) -> str:
        prompt = f'''Please grade the following student solution based on the exercise description and grading rubric provided.

        Exercise Description:
        {exercise}

        Grading Rubric:
        {rubric}

        Students Solution:
        {solution}

        Please provide:
        1. A detailed analysis of the solution
        2. Points awarded for each rubric criterion
        3. Total score
        4. Specific feedback and suggestions for improvement

        Format your response ONLY as JSON with the following structure:
        {{
            "analysis": "detailed analysis here",
            "criterion_scores": {{
                "criterion1": "score",
                "criterion2": "score"
            }},
            "total_score": "numerical_score",
            "feedback": "detailed feedback here"
        }}'''
        return prompt

    def grade_solution(self, exercise_path: str, rubric_path: str, solution_path: str) -> Dict:
        try:
            exercise_text = self.read_document(exercise_path)
            rubric_text = self.read_document(rubric_path)
            solution_text = self.read_document(solution_path)

            prompt = self.prepare_grading_prompt(exercise_text, rubric_text, solution_text)

            response = self.client.chat.completions.create(
                model="gpt-4-turbo-preview",
                messages=[
                    {"role": "system", "content": "You are an expert grader who evaluates student solutions accurately and fairly."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.0,
                max_tokens=4000,
                response_format={"type": "json_object"}
            )

            response_content = response.choices[0].message.content.strip()
            print("Raw API Response:", response_content)

            try:
                grading_result = json.loads(response_content)
                return grading_result
            except json.JSONDecodeError:
                raise Exception("Failed to parse API response as JSON")

        except Exception as e:
            raise Exception(f"Grading failed: {str(e)}")

    def save_results(self, results: Dict, output_path: str):
        try:
            doc = docx.Document()
            doc.add_heading('Grading Results', 0)
            
            doc.add_heading('Analysis', level=1)
            doc.add_paragraph(str(results.get('analysis', 'No analysis provided')))

            doc.add_heading('Criterion Scores', level=1)
            criterion_scores = results.get('criterion_scores', {})
            if isinstance(criterion_scores, dict):
                for criterion, score in criterion_scores.items():
                    doc.add_paragraph(f"{criterion}: {score}")
            else:
                doc.add_paragraph("No criterion scores available")
            
            doc.add_heading('Total Score', level=1)
            doc.add_paragraph(str(results.get('total_score', 'No score available')))
            
            doc.add_heading('Feedback', level=1)
            doc.add_paragraph(str(results.get('feedback', 'No feedback provided')))
            
            doc.save(output_path)
            print(f"Results successfully saved to {output_path}")

        except Exception as e:
            raise Exception(f"Error saving results: {str(e)}")

def main():
    try:
        MI_CLAVE = os.getenv('MI_CLAVE_API_OPENAI')
        if not MI_CLAVE:
            raise ValueError("OpenAI API key not found in environment variables")

        grader = DocumentGrader(MI_CLAVE)

        # Fixed paths for exercise and rubric
        exercise_path = "/Users/JJVR/Library/CloudStorage/OneDrive-KALEIDAGEOGRAFIAS&MERCADOSSL/ESBS/Asignaturas/Plagiarism & AI/Exercise/exercise instructions.md"
        rubric_path = "/Users/JJVR/Library/CloudStorage/OneDrive-KALEIDAGEOGRAFIAS&MERCADOSSL/ESBS/Asignaturas/Plagiarism & AI/Exercise/grading rubric.md"

        # Select folder with student solutions
        print("Please select the folder containing student solutions...")
        solutions_folder = select_folder()
        if not solutions_folder:
            raise ValueError("No folder was selected")

        # Select output folder
        print("Please select where to save the grading results...")
        output_folder = select_output_folder()
        if not output_folder:
            raise ValueError("No output folder was selected")

        # Verify exercise and rubric files exist
        for path, description in [
            (exercise_path, "Exercise instructions"),
            (rubric_path, "Grading rubric")
        ]:
            if not os.path.exists(path):
                raise FileNotFoundError(f"{description} file not found: {path}")

        # Process all files in the selected folder
        supported_extensions = {'.pdf', '.docx'}
        for filename in os.listdir(solutions_folder):
            if os.path.splitext(filename)[1].lower() in supported_extensions:
                solution_path = os.path.join(solutions_folder, filename)
                output_filename = f"grading_results_{os.path.splitext(filename)[0]}_graded.docx"
                output_path = os.path.join(output_folder, output_filename)

                print(f"Grading solution: {filename}")
                try:
                    results = grader.grade_solution(exercise_path, rubric_path, solution_path)
                    grader.save_results(results, output_path)
                    print(f"Grading completed successfully for {filename}")
                except Exception as e:
                    print(f"Error grading {filename}: {str(e)}")
                    continue

        print("All grading completed!")

    except Exception as e:
        print(f"Error: {str(e)}")
        exit(1)

if __name__ == "__main__":
    main()
